#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, font_name: str, font_size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, font_name: str = "Microsoft YaHei", font_size: int = 10, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    set_run_font(run, font_name, font_size, bold)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").replace(" ", "")
    return bool(stripped) and all(ch == "-" or ch == ":" or ch == "|" for ch in stripped)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if is_table_separator(line):
            continue
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))
    doc.add_paragraph("")


def add_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    set_run_font(run, "Microsoft YaHei", 11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text.strip())
    set_run_font(run, "Microsoft YaHei", 11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text.strip())
    set_run_font(run, "Microsoft YaHei", 11)


def add_heading(doc: Document, level: int, text: str) -> None:
    if level <= 1:
        heading = doc.add_paragraph(style="Title")
        run = heading.add_run(text.strip())
        set_run_font(run, "Microsoft YaHei", 22, True)
        heading.paragraph_format.space_after = Pt(10)
        return
    word_level = min(level - 1, 9)
    p = doc.add_heading(level=word_level)
    run = p.add_run(text.strip())
    size = 16 if word_level == 1 else 14 if word_level == 2 else 12
    set_run_font(run, "Microsoft YaHei", size, True)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = 1
    run = header.add_run("图书多渠道商品中台 | 一期产品说明（客户版）")
    set_run_font(run, "Microsoft YaHei", 9)

    footer = section.footer.paragraphs[0]
    footer.alignment = 1
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def render_markdown(md_path: Path, output_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    add_header_footer(doc)

    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph_buffer() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph(doc, " ".join(line.strip() for line in paragraph_buffer).replace("**", ""))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph_buffer()
            i += 1
            continue

        if stripped == "---":
            flush_paragraph_buffer()
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph_buffer()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph_buffer()
            level = len(heading_match.group(1))
            text_value = heading_match.group(2).replace("**", "")
            add_heading(doc, level, text_value)
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph_buffer()
            add_bullet(doc, bullet_match.group(1).replace("**", ""))
            i += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            flush_paragraph_buffer()
            add_numbered(doc, numbered_match.group(1).replace("**", ""))
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph_buffer()
    doc.save(output_path)


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: render_customer_docx.py <input.md> <output.docx>")
        return 1
    md_path = Path(argv[1]).resolve()
    output_path = Path(argv[2]).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    render_markdown(md_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
